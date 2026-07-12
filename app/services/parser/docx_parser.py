"""
app/services/parser/docx_parser.py — DOCX → raw text extractor.

Uses python-docx. Preserves paragraph order and table cell text.
Raises ParseError if the file is not a valid DOCX or produces no usable text.
"""
import io
import logging

logger = logging.getLogger(__name__)


class ParseError(Exception):
    """Raised when a file cannot be parsed into usable text."""
    pass


def extract_text_from_docx(file_bytes: bytes, filename: str = "unknown.docx") -> str:
    """
    Extract raw text from a DOCX file.

    Args:
        file_bytes: Raw bytes of the DOCX file.
        filename: Original filename (for logging).

    Returns:
        Extracted text as a string with paragraph breaks.

    Raises:
        ParseError: If python-docx cannot open the file or text is insufficient.
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(io.BytesIO(file_bytes))
        sections: list[str] = []

        # --- Main body paragraphs ---
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                sections.append(text)

        # --- Table cells (skills tables, education grids, etc.) ---
        for table in doc.tables:
            for row in table.rows:
                row_texts: list[str] = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    sections.append(" | ".join(row_texts))

        # --- Text boxes and headers/footers (via XML) ---
        try:
            body = doc.element.body
            for textbox in body.iter(qn("w:txbxContent")):
                for para in textbox.iter(qn("w:p")):
                    text = "".join(
                        r.text for r in para.iter(qn("w:t")) if r.text
                    ).strip()
                    if text:
                        sections.append(text)
        except Exception as xml_exc:
            logger.debug("Text box extraction skipped: %s", xml_exc)

        result = "\n".join(sections)

        if len(result.strip()) < 50:
            raise ParseError(
                f"Could not extract usable text from '{filename}'. "
                "The file may be empty, corrupted, or contain only images."
            )

        logger.info("python-docx extracted %d chars from '%s'", len(result), filename)
        return result

    except ParseError:
        raise
    except Exception as exc:
        raise ParseError(
            f"Failed to open '{filename}' as a DOCX file: {exc}. "
            "Ensure the file is a valid .docx (not .doc or a renamed PDF)."
        ) from exc
