"""
tests/generate_sample_resumes.py
Generates 3 realistic synthetic resume files for Phase 1 testing.

Run with:  py -3 tests/generate_sample_resumes.py

Produces:
  tests/sample_resumes/resume_backend_engineer.pdf
  tests/sample_resumes/resume_data_scientist.docx
  tests/sample_resumes/resume_fullstack_dev.pdf
"""
import sys
from pathlib import Path

OUTPUT_DIR = Path(__file__).parent / "sample_resumes"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Resume 1: Backend Engineer (PDF via reportlab)
# ---------------------------------------------------------------------------
BACKEND_TEXT = """\
Aisha Raza
aisha.raza@email.com | +1 (312) 555-0178 | linkedin.com/in/aisharaza | github.com/aisharaza

PROFESSIONAL SUMMARY
Backend engineer with 5 years of experience building scalable Python microservices,
REST APIs, and data pipelines. Strong focus on system design, observability, and
developer tooling.

TECHNICAL SKILLS
Languages: Python, Go, SQL, Bash
Frameworks: FastAPI, Django, Flask, gRPC
Databases: PostgreSQL, Redis, MongoDB, Elasticsearch
Infrastructure: Docker, Kubernetes, Terraform, AWS (EC2, RDS, S3, Lambda)
Tools: Git, GitHub Actions, Prometheus, Grafana, Celery, RabbitMQ
Testing: pytest, unittest, Postman

EXPERIENCE

Senior Backend Engineer — DataStream Inc., Chicago, IL
Jan 2022 – Present
- Designed and deployed a high-throughput event ingestion service (FastAPI + Kafka)
  processing 2M+ events/day with <50ms p99 latency
- Migrated monolith to microservices, reducing deployment time by 60%
- Built internal developer CLI tooling used by 30+ engineers
- Introduced structured logging with OpenTelemetry; reduced MTTR by 40%

Backend Engineer — CloudOps Solutions, Chicago, IL
Jun 2019 – Dec 2021
- Built RESTful APIs using Django REST Framework serving 500K daily active users
- Optimized PostgreSQL queries (added indexes, rewrote N+1s); avg query time -70%
- Implemented background job processing with Celery + Redis
- Wrote integration tests achieving 85% code coverage

EDUCATION

Bachelor of Science in Computer Science
University of Illinois at Chicago — 2019

CERTIFICATIONS
AWS Certified Developer – Associate (2022)
Docker Certified Associate (2021)

PROJECTS

EventBridge CLI (github.com/aisharaza/eventbridge-cli)
A developer tool for testing and replaying Kafka events locally.
Built with Python, Click, Kafka, Docker

PGOptimizer
Automated PostgreSQL slow query analyzer with actionable recommendations.
Technologies: Python, psycopg2, PostgreSQL, Flask
"""


def generate_backend_pdf():
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_LEFT

        out_path = OUTPUT_DIR / "resume_backend_engineer.pdf"
        doc = SimpleDocTemplate(
            str(out_path),
            pagesize=letter,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )
        styles = getSampleStyleSheet()
        name_style = ParagraphStyle("Name", fontSize=16, fontName="Helvetica-Bold", spaceAfter=4)
        contact_style = ParagraphStyle("Contact", fontSize=9, fontName="Helvetica", spaceAfter=8)
        section_style = ParagraphStyle(
            "Section", fontSize=11, fontName="Helvetica-Bold",
            textColor=colors.HexColor("#1a1a2e"), spaceBefore=10, spaceAfter=4,
            borderPad=2,
        )
        body_style = ParagraphStyle("Body", fontSize=9, fontName="Helvetica", leading=13, spaceAfter=2)

        story = []
        for line in BACKEND_TEXT.strip().split("\n"):
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 4))
            elif stripped == line.strip() and stripped.isupper() and len(stripped) < 40:
                story.append(Paragraph(stripped, section_style))
            elif line.startswith("  ") or line.startswith("\t"):
                story.append(Paragraph(f"&bull; {stripped.lstrip('- ')}", body_style))
            else:
                story.append(Paragraph(stripped, body_style))

        doc.build(story)
        print(f"✓ Generated: {out_path}")
        return True

    except ImportError:
        # Fallback: plain text PDF-like file using fpdf2
        try:
            from fpdf import FPDF
            out_path = OUTPUT_DIR / "resume_backend_engineer.pdf"
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", size=10)
            for line in BACKEND_TEXT.split("\n"):
                pdf.cell(0, 5, txt=line[:100], ln=True)
            pdf.output(str(out_path))
            print(f"✓ Generated (fpdf2): {out_path}")
            return True
        except ImportError:
            print("✗ Neither reportlab nor fpdf2 installed. Install with: pip install reportlab")
            # Write as a .txt instead so tests can still run
            out_path = OUTPUT_DIR / "resume_backend_engineer.txt"
            out_path.write_text(BACKEND_TEXT, encoding="utf-8")
            print(f"  Wrote plain text fallback: {out_path}")
            return False


# ---------------------------------------------------------------------------
# Resume 2: Data Scientist (DOCX via python-docx)
# ---------------------------------------------------------------------------
DATA_SCIENTIST_DATA = {
    "name": "Marcus Chen",
    "contact": "mchen@datamail.io | +1 (415) 555-0234 | linkedin.com/in/marcuschen | github.com/marcuschen-ds",
    "summary": (
        "Data Scientist with 4 years of experience in machine learning, NLP, and "
        "statistical modeling. Delivered production ML models used by 2M+ users. "
        "PhD candidate in Computational Statistics, UC Berkeley."
    ),
    "skills": [
        ("Languages", "Python, R, SQL, Julia"),
        ("ML/DL", "scikit-learn, PyTorch, TensorFlow, XGBoost, LightGBM, Hugging Face Transformers"),
        ("NLP", "spaCy, NLTK, Gensim, sentence-transformers, BM25, TF-IDF"),
        ("Data", "pandas, NumPy, Dask, Apache Spark, Airflow"),
        ("Visualization", "matplotlib, seaborn, Plotly, Tableau"),
        ("Infrastructure", "Docker, AWS SageMaker, MLflow, DVC, PostgreSQL"),
    ],
    "experience": [
        {
            "title": "Senior Data Scientist",
            "company": "HealthAI Corp",
            "location": "San Francisco, CA",
            "duration": "Mar 2022 – Present",
            "bullets": [
                "Built NLP pipeline (spaCy + BERT) to extract medical entities from clinical notes — 91% F1 score",
                "Developed patient readmission risk model (XGBoost); reduced ICU readmission rate by 18%",
                "Productionized 6 ML models on SageMaker serving 50K daily predictions",
                "Mentored 3 junior data scientists; established MLOps best practices",
            ],
        },
        {
            "title": "Data Scientist",
            "company": "RetailSense Inc.",
            "location": "Oakland, CA",
            "duration": "Aug 2020 – Feb 2022",
            "bullets": [
                "Built recommendation engine (collaborative filtering + content-based) increasing CTR by 23%",
                "Designed A/B testing framework used across 15 product experiments",
                "Created real-time demand forecasting pipeline with ARIMA + Prophet",
            ],
        },
    ],
    "education": [
        ("PhD Candidate, Computational Statistics", "UC Berkeley", "2020 – Present"),
        ("Bachelor of Science in Statistics", "Stanford University", "2020"),
    ],
    "projects": [
        ("ResumeMatcher", "NLP-based resume-job matching using TF-IDF + BM25 scoring", "Python, spaCy, FastAPI"),
        ("ClinicalNLP", "Named entity recognition for clinical text (open source)", "Python, Hugging Face, PyTorch"),
    ],
    "certifications": [
        "AWS Certified Machine Learning – Specialty (2023)",
        "Google Professional Data Engineer (2022)",
        "Deep Learning Specialization – Coursera (2021)",
    ],
}


def generate_data_scientist_docx():
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        out_path = OUTPUT_DIR / "resume_data_scientist.docx"
        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        def add_heading(text, level=1):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14 if level == 1 else 11)
            if level == 1:
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            p.paragraph_format.space_before = Pt(10 if level > 1 else 0)
            p.paragraph_format.space_after = Pt(4)
            return p

        def add_body(text, bold=False, italic=False):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.bold = bold
            run.italic = italic
            p.paragraph_format.space_after = Pt(2)
            return p

        def add_bullet(text):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(text)
            run.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(1)

        d = DATA_SCIENTIST_DATA

        # Name
        name_p = doc.add_paragraph()
        name_run = name_p.add_run(d["name"])
        name_run.bold = True
        name_run.font.size = Pt(16)
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact
        contact_p = doc.add_paragraph()
        contact_p.add_run(d["contact"]).font.size = Pt(9)
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Summary
        add_heading("PROFESSIONAL SUMMARY", level=2)
        add_body(d["summary"])

        # Skills
        add_heading("TECHNICAL SKILLS", level=2)
        for category, skills in d["skills"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{category}: ")
            run.bold = True
            run.font.size = Pt(9)
            p.add_run(skills).font.size = Pt(9)
            p.paragraph_format.space_after = Pt(2)

        # Experience
        add_heading("EXPERIENCE", level=2)
        for exp in d["experience"]:
            p = doc.add_paragraph()
            r1 = p.add_run(f"{exp['title']} — {exp['company']}, {exp['location']}")
            r1.bold = True
            r1.font.size = Pt(10)
            add_body(exp["duration"], italic=True)
            for bullet in exp["bullets"]:
                add_bullet(bullet)

        # Education
        add_heading("EDUCATION", level=2)
        for degree, school, years in d["education"]:
            p = doc.add_paragraph()
            p.add_run(f"{degree}\n").bold = True
            p.runs[-1].font.size = Pt(9)
            p.add_run(f"{school} — {years}").font.size = Pt(9)

        # Projects
        add_heading("PROJECTS", level=2)
        for proj_name, proj_desc, proj_tech in d["projects"]:
            p = doc.add_paragraph()
            r = p.add_run(f"{proj_name}: ")
            r.bold = True
            r.font.size = Pt(9)
            p.add_run(f"{proj_desc} | Technologies: {proj_tech}").font.size = Pt(9)

        # Certifications
        add_heading("CERTIFICATIONS", level=2)
        for cert in d["certifications"]:
            add_bullet(cert)

        doc.save(str(out_path))
        print(f"✓ Generated: {out_path}")
        return True

    except ImportError:
        print("✗ python-docx not installed. Install with: pip install python-docx")
        out_path = OUTPUT_DIR / "resume_data_scientist.txt"
        out_path.write_text(
            "\n".join([
                DATA_SCIENTIST_DATA["name"],
                DATA_SCIENTIST_DATA["contact"],
                "\nSUMMARY\n" + DATA_SCIENTIST_DATA["summary"],
            ]),
            encoding="utf-8",
        )
        print(f"  Wrote plain text fallback: {out_path}")
        return False


# ---------------------------------------------------------------------------
# Resume 3: Full-Stack Developer (PDF)
# ---------------------------------------------------------------------------
FULLSTACK_TEXT = """\
Priya Nair
priya.nair@dev.io | +44 7700 900456 | github.com/priyanair | linkedin.com/in/priyanair-dev

SUMMARY
Full-stack engineer with 6 years building React frontends and Node.js/Python backends.
Led migration of legacy Angular app to React, improving Lighthouse score from 54 to 91.
Passionate about accessibility, performance, and design systems.

SKILLS
Frontend: JavaScript, TypeScript, React, Next.js, Vue.js, HTML5, CSS3, Tailwind CSS, Webpack
Backend: Node.js, Express, Python, FastAPI, GraphQL, REST APIs
Databases: PostgreSQL, MySQL, MongoDB, Redis, Firebase
DevOps: Docker, AWS (CloudFront, Lambda, DynamoDB), Vercel, Netlify, GitHub Actions
Testing: Jest, React Testing Library, Cypress, Playwright
Design: Figma, Storybook, Accessible design (WCAG 2.1)

EXPERIENCE

Lead Frontend Engineer — FinFlow Ltd., London, UK
Feb 2021 – Present
- Led migration from Angular to React + Next.js (SSR); Lighthouse score 54 → 91
- Built company-wide design system (50+ components) using Storybook + Tailwind CSS
- Implemented real-time dashboard with WebSockets serving 10K concurrent users
- Reduced bundle size by 45% via code splitting and lazy loading
- Managed team of 4 frontend engineers; conducted code reviews and weekly 1:1s

Full-Stack Developer — Nexio Digital, Manchester, UK
Sep 2018 – Jan 2021
- Built customer portal (React + Node.js + PostgreSQL) from scratch, 0 → 25K users
- Developed GraphQL API replacing 12 REST endpoints; 30% reduction in over-fetching
- Integrated Stripe payment processing for subscription billing
- Set up CI/CD pipelines with GitHub Actions and AWS CodeDeploy

EDUCATION

Bachelor of Engineering in Software Engineering
University of Manchester — 2018

PROJECTS

OpenDash (github.com/priyanair/opendash)
Open-source analytics dashboard built with Next.js, Recharts, and PostgreSQL.
Used by 200+ developers on GitHub. Technologies: TypeScript, Next.js, PostgreSQL, Tailwind CSS

AccessKit
A11y testing toolkit for React component libraries.
Built with React, Jest, axe-core, Storybook

CERTIFICATIONS
AWS Certified Cloud Practitioner (2022)
Google UX Design Certificate (2021)
Meta Front-End Developer Certificate (2020)
"""


def generate_fullstack_pdf():
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

        out_path = OUTPUT_DIR / "resume_fullstack_dev.pdf"
        doc = SimpleDocTemplate(
            str(out_path),
            pagesize=letter,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        name_style = ParagraphStyle("Name", fontSize=16, fontName="Helvetica-Bold", spaceAfter=2)
        contact_style = ParagraphStyle("Contact", fontSize=9, fontName="Helvetica", spaceAfter=6)
        section_style = ParagraphStyle(
            "Section", fontSize=11, fontName="Helvetica-Bold",
            textColor=colors.HexColor("#2d3436"), spaceBefore=8, spaceAfter=3,
        )
        job_title_style = ParagraphStyle(
            "JobTitle", fontSize=10, fontName="Helvetica-Bold", spaceAfter=1
        )
        body_style = ParagraphStyle("Body", fontSize=9, fontName="Helvetica", leading=13, spaceAfter=2)
        bullet_style = ParagraphStyle(
            "Bullet", fontSize=9, fontName="Helvetica", leading=13,
            leftIndent=12, spaceAfter=1,
        )

        story = []
        lines = FULLSTACK_TEXT.strip().split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if i == 0:
                story.append(Paragraph(stripped, name_style))
            elif i == 1:
                story.append(Paragraph(stripped, contact_style))
                story.append(HRFlowable(width="100%", thickness=1, color=colors.lightgrey))
            elif not stripped:
                story.append(Spacer(1, 3))
            elif stripped.isupper() and len(stripped) < 50 and len(stripped) > 2:
                story.append(Paragraph(stripped, section_style))
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
            elif stripped.startswith("- "):
                story.append(Paragraph(f"• {stripped[2:]}", bullet_style))
            else:
                story.append(Paragraph(stripped, body_style))
            i += 1

        doc.build(story)
        print(f"✓ Generated: {out_path}")
        return True

    except ImportError:
        try:
            from fpdf import FPDF
            out_path = OUTPUT_DIR / "resume_fullstack_dev.pdf"
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", size=10)
            for line in FULLSTACK_TEXT.split("\n"):
                pdf.cell(0, 5, txt=line[:100], ln=True)
            pdf.output(str(out_path))
            print(f"✓ Generated (fpdf2): {out_path}")
            return True
        except ImportError:
            out_path = OUTPUT_DIR / "resume_fullstack_dev.txt"
            out_path.write_text(FULLSTACK_TEXT, encoding="utf-8")
            print(f"  Wrote plain text fallback: {out_path}")
            return False


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    print(f"\nGenerating sample resumes -> {OUTPUT_DIR}\n")
    results = [
        generate_backend_pdf(),
        generate_data_scientist_docx(),
        generate_fullstack_pdf(),
    ]
    success = sum(results)
    print(f"\n{'='*50}")
    print(f"Generated {success}/3 sample resume files in {OUTPUT_DIR}")
    if success < 3:
        print("\nSome files fell back to plain text. Install missing deps:")
        print("  pip install reportlab python-docx")
    sys.exit(0 if success > 0 else 1)
